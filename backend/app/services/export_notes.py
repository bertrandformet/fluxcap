import io
import re

from docx import Document

from app.models import Note

_RE_GRAS_ITALIQUE = re.compile(r"(\*\*.+?\*\*|\*.+?\*)")


def _corps_markdown(note: Note) -> str:
    return note.contenu or note.apercu or ""


def generer_txt(note: Note) -> bytes:
    corps = _corps_markdown(note)
    return f"{note.titre}\n\n{corps}".encode("utf-8")


def generer_md(note: Note) -> bytes:
    lignes = [f"# {note.titre}", ""]
    if note.url:
        lignes.append(f"<{note.url}>")
        lignes.append("")
    lignes.append(_corps_markdown(note))
    return "\n".join(lignes).encode("utf-8")


def _ajouter_texte_avec_style(paragraphe, texte: str) -> None:
    for morceau in _RE_GRAS_ITALIQUE.split(texte):
        if not morceau:
            continue
        if morceau.startswith("**") and morceau.endswith("**"):
            paragraphe.add_run(morceau[2:-2]).bold = True
        elif morceau.startswith("*") and morceau.endswith("*"):
            paragraphe.add_run(morceau[1:-1]).italic = True
        else:
            paragraphe.add_run(morceau)


def generer_docx(note: Note) -> bytes:
    """Conversion markdown → docx minimale : titres, listes, citations, gras/italique."""
    doc = Document()
    doc.add_heading(note.titre, level=0)

    if note.url:
        doc.add_paragraph().add_run(note.url).italic = True

    for ligne in _corps_markdown(note).splitlines():
        brute = ligne.rstrip()
        if not brute.strip():
            doc.add_paragraph("")
            continue

        entete = re.match(r"^(#{1,5})\s+(.*)", brute)
        if entete:
            doc.add_heading(entete.group(2), level=min(len(entete.group(1)), 4))
            continue

        if re.match(r"^[-*]\s+", brute):
            p = doc.add_paragraph(style="List Bullet")
            _ajouter_texte_avec_style(p, re.sub(r"^[-*]\s+", "", brute))
            continue

        numerote = re.match(r"^\d+\.\s+(.*)", brute)
        if numerote:
            p = doc.add_paragraph(style="List Number")
            _ajouter_texte_avec_style(p, numerote.group(1))
            continue

        if brute.startswith(">"):
            p = doc.add_paragraph(style="Intense Quote")
            _ajouter_texte_avec_style(p, brute.lstrip(">").strip())
            continue

        p = doc.add_paragraph()
        _ajouter_texte_avec_style(p, brute)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


GENERATEURS = {
    "txt": (generer_txt, "text/plain"),
    "md": (generer_md, "text/markdown"),
    "docx": (generer_docx, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
}
